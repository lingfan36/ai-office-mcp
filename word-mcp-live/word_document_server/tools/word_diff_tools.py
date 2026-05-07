"""Word document diff tool — compare two .docx files (cross-platform, no Office needed)."""
import json
from pathlib import Path


def word_diff(
    path1: str,
    path2: str,
    max_diffs: int = 200,
) -> str:
    """Compare two Word documents and return paragraph- and table-level differences.

    Works cross-platform with no Word installation required.

    Args:
        path1: Path to the "before" document.
        path2: Path to the "after" document.
        max_diffs: Cap on returned differences (avoids huge payloads).

    Returns JSON with:
        - changed_paragraphs: list of {index, before, after, style_before, style_after}
        - added_paragraphs: paragraphs present only in path2 (by index position)
        - removed_paragraphs: paragraphs present only in path1 (by index position)
        - changed_tables: list of {table_index, row, col, before, after}
        - truncated: true if max_diffs was hit
        - summary: counts
    """
    try:
        import docx

        p1, p2 = Path(path1).resolve(), Path(path2).resolve()
        if not p1.exists():
            return json.dumps({"success": False, "error": f"File not found: {path1}"})
        if not p2.exists():
            return json.dumps({"success": False, "error": f"File not found: {path2}"})

        doc1 = docx.Document(str(p1))
        doc2 = docx.Document(str(p2))

        paras1 = [(p.text, p.style.name) for p in doc1.paragraphs]
        paras2 = [(p.text, p.style.name) for p in doc2.paragraphs]

        changed_paragraphs = []
        truncated = False
        min_len = min(len(paras1), len(paras2))

        for i in range(min_len):
            text1, style1 = paras1[i]
            text2, style2 = paras2[i]
            if text1 != text2 or style1 != style2:
                changed_paragraphs.append({
                    "index": i,
                    "before": text1,
                    "after": text2,
                    "style_before": style1,
                    "style_after": style2,
                    "style_changed": style1 != style2,
                })
                if len(changed_paragraphs) >= max_diffs:
                    truncated = True
                    break

        removed_paragraphs = []
        added_paragraphs = []
        if not truncated:
            if len(paras1) > min_len:
                removed_paragraphs = [
                    {"index": i, "text": paras1[i][0], "style": paras1[i][1]}
                    for i in range(min_len, len(paras1))
                ]
            elif len(paras2) > min_len:
                added_paragraphs = [
                    {"index": i, "text": paras2[i][0], "style": paras2[i][1]}
                    for i in range(min_len, len(paras2))
                ]

        # Compare tables
        changed_tables = []
        tables1 = doc1.tables
        tables2 = doc2.tables
        added_tables = max(0, len(tables2) - len(tables1))
        removed_tables = max(0, len(tables1) - len(tables2))

        if not truncated:
            for t_idx in range(min(len(tables1), len(tables2))):
                tbl1, tbl2 = tables1[t_idx], tables2[t_idx]
                min_rows = min(len(tbl1.rows), len(tbl2.rows))
                done = False
                for r_idx in range(min_rows):
                    cells1 = tbl1.rows[r_idx].cells
                    cells2 = tbl2.rows[r_idx].cells
                    min_cols = min(len(cells1), len(cells2))
                    for c_idx in range(min_cols):
                        v1 = cells1[c_idx].text
                        v2 = cells2[c_idx].text
                        if v1 != v2:
                            changed_tables.append({
                                "table_index": t_idx,
                                "row": r_idx,
                                "col": c_idx,
                                "before": v1,
                                "after": v2,
                            })
                            if len(changed_tables) >= max_diffs:
                                truncated = True
                                done = True
                                break
                    if done:
                        break
                if done:
                    break

        return json.dumps({
            "success": True,
            "path1": str(p1),
            "path2": str(p2),
            "changed_paragraphs": changed_paragraphs,
            "added_paragraphs": added_paragraphs,
            "removed_paragraphs": removed_paragraphs,
            "changed_tables": changed_tables,
            "truncated": truncated,
            "summary": {
                "changed_paragraphs": len(changed_paragraphs),
                "added_paragraphs": len(added_paragraphs),
                "removed_paragraphs": len(removed_paragraphs),
                "changed_tables": len(changed_tables),
                "added_tables": added_tables,
                "removed_tables": removed_tables,
                "paragraphs_doc1": len(paras1),
                "paragraphs_doc2": len(paras2),
                "tables_doc1": len(tables1),
                "tables_doc2": len(tables2),
            },
        })
    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})
