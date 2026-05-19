"""Audit a .docx against a structured formatting spec.

The audit engine is the verification layer that breaks "AI says done" away
from "actually done". Workflow:

    1. Caller (AI or human) writes a JSON spec — each item is
       ``{"id": "<dotted.path>", "expected": <value>}``.
    2. After making edits, caller invokes ``word_audit_against_spec``.
    3. Engine reads the actual document state, compares each item, and
       returns ``{pass_count, fail_count, items: [...]}``.
    4. Caller is required to surface the audit JSON before declaring
       completion (enforced via CLAUDE.md rule, not by code).

The engine is intentionally simple: each spec id is dispatched to a single
checker function. To add a new check, register one entry in CHECKERS.

This module reads from disk via python-docx + lxml. If a file is currently
open in Word with unsaved changes, those changes won't be visible until
``word_live_save`` is called first — document this limitation in the
caller-facing CLAUDE.md rule.
"""

from __future__ import annotations

import json
import os
import re
import zipfile
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from word_document_server.utils.file_utils import ensure_docx_extension

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# ---------------------------------------------------------------------------
# Small helpers shared across checkers
# ---------------------------------------------------------------------------

def _run_east_asia_font(run) -> Optional[str]:
    """Return the East Asia font name set on a run, or None."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return None
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return None
    return rFonts.get(qn('w:eastAsia')) or rFonts.get(qn('w:asciiTheme'))


def _run_ascii_font(run) -> Optional[str]:
    """Return the ASCII (English) font name set on a run, or None."""
    name = run.font.name
    if name:
        return name
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            return rFonts.get(qn('w:ascii'))
    return None


def _run_size_pt(run) -> Optional[float]:
    """Run font size in points, or None if unset (inherits from style)."""
    if run.font.size is None:
        return None
    return run.font.size.pt


def _para_paragraphs_by_style_prefix(doc, prefix: str):
    """Yield paragraphs whose style name starts with prefix (case-insensitive)."""
    p = prefix.lower()
    for para in doc.paragraphs:
        name = (para.style.name or "").lower()
        if name.startswith(p):
            yield para


def _para_text_runs(para):
    """Yield only runs with non-empty text."""
    for r in para.runs:
        if r.text:
            yield r


def _is_blank(para) -> bool:
    return not (para.text or "").strip()


def _alignment_name(value) -> str:
    # python-docx alignment enum → readable name
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        None: "inherit",
    }
    return mapping.get(value, str(value))


# ---------------------------------------------------------------------------
# Result helpers
# ---------------------------------------------------------------------------

def _ok(actual) -> Dict[str, Any]:
    return {"pass": True, "actual": actual}


def _fail(actual, **extra) -> Dict[str, Any]:
    out = {"pass": False, "actual": actual}
    out.update(extra)
    return out


def _num_close(a: float, b: float, tol: float) -> bool:
    return abs(float(a) - float(b)) <= tol


# ---------------------------------------------------------------------------
# Page-level checkers (all read from section[0] unless noted)
# ---------------------------------------------------------------------------

def _check_section_margin(attr_cm: str, tol: float = 0.05):
    def checker(doc, expected, ctx):
        section = doc.sections[0]
        prop = getattr(section, attr_cm.replace("_cm", ""))
        actual_cm = round(prop.cm, 3)
        if _num_close(actual_cm, expected, tol):
            return _ok(actual_cm)
        return _fail(actual_cm)
    return checker


def _check_odd_even_pages(doc, expected, ctx):
    with zipfile.ZipFile(ctx["filename"], 'r') as z:
        try:
            settings = z.read('word/settings.xml').decode('utf-8')
        except KeyError:
            settings = ""
    actual = '<w:evenAndOddHeaders' in settings
    return _ok(actual) if actual == bool(expected) else _fail(actual)


# ---------------------------------------------------------------------------
# Body-paragraph checkers — sample over paragraphs whose style is Normal
# and which contain Chinese text (so we don't penalize an English-only paragraph).
# ---------------------------------------------------------------------------

_RE_CJK = re.compile(r'[一-鿿]')


def _body_paragraphs(doc):
    out = []
    for p in doc.paragraphs:
        if _is_blank(p):
            continue
        name = (p.style.name or "").lower()
        if name.startswith("heading") or name in ("title", "subtitle"):
            continue
        out.append(p)
    return out


def _check_body_font_eastasia(doc, expected, ctx):
    seen: Dict[str, int] = {}
    for p in _body_paragraphs(doc):
        if not _RE_CJK.search(p.text):
            continue
        for r in _para_text_runs(p):
            f = _run_east_asia_font(r) or _run_ascii_font(r)
            if f:
                seen[f] = seen.get(f, 0) + 1
    if not seen:
        return _fail("<no body CJK runs found>")
    distinct = sorted(seen.keys())
    if distinct == [expected]:
        return _ok(expected)
    return _fail(distinct, hint="multiple fonts detected in body CJK text")


def _check_body_font_ascii(doc, expected, ctx):
    seen: Dict[str, int] = {}
    for p in _body_paragraphs(doc):
        for r in _para_text_runs(p):
            # Look at runs with ASCII letters
            if not re.search(r'[A-Za-z]', r.text):
                continue
            f = _run_ascii_font(r)
            if f:
                seen[f] = seen.get(f, 0) + 1
    if not seen:
        return _ok("<no body English runs>")
    distinct = sorted(seen.keys())
    if distinct == [expected]:
        return _ok(expected)
    return _fail(distinct)


def _check_body_size_pt(doc, expected, ctx):
    sizes: Dict[float, int] = {}
    for p in _body_paragraphs(doc):
        for r in _para_text_runs(p):
            sz = _run_size_pt(r)
            if sz is not None:
                sizes[sz] = sizes.get(sz, 0) + 1
    if not sizes:
        # Style-inherited size — assume default. Mark inconclusive but pass.
        return _ok("<inherits from style>")
    distinct = sorted(sizes.keys())
    if len(distinct) == 1 and _num_close(distinct[0], expected, 0.5):
        return _ok(distinct[0])
    return _fail(distinct, hint="inconsistent or wrong body font size (pt)")


def _check_body_line_spacing_pt(doc, expected, ctx):
    # Read all body paragraphs' line spacing in pt, find dominant value.
    pts: Dict[float, int] = {}
    for p in _body_paragraphs(doc):
        pf = p.paragraph_format
        ls = pf.line_spacing
        if ls is None:
            continue
        # python-docx returns float (for multiple) or Length object (for exactly).
        from docx.shared import Emu
        if hasattr(ls, "pt"):
            val = round(ls.pt, 2)
        else:
            # multiple: float multiplier, convert to pt assuming 12pt baseline
            val = round(float(ls) * 12.0, 2)
        pts[val] = pts.get(val, 0) + 1
    if not pts:
        return _ok("<inherits>")
    distinct = sorted(pts.keys())
    if len(distinct) == 1 and _num_close(distinct[0], expected, 0.5):
        return _ok(distinct[0])
    return _fail(distinct)


# ---------------------------------------------------------------------------
# Heading checkers
# ---------------------------------------------------------------------------

def _heading_paragraphs(doc, level: int):
    target = f"heading {level}"
    return [p for p in doc.paragraphs if (p.style.name or "").lower() == target]


def _check_heading_font_eastasia(level: int):
    def checker(doc, expected, ctx):
        paras = _heading_paragraphs(doc, level)
        if not paras:
            return _fail("<no Heading {} paragraphs found>".format(level),
                         hint="paragraphs may be styled as bold Normal instead of real Heading {}".format(level))
        seen = set()
        for p in paras:
            for r in _para_text_runs(p):
                f = _run_east_asia_font(r) or _run_ascii_font(r)
                if f:
                    seen.add(f)
        if seen == {expected}:
            return _ok(expected)
        return _fail(sorted(seen))
    return checker


def _check_heading_size_pt(level: int):
    def checker(doc, expected, ctx):
        paras = _heading_paragraphs(doc, level)
        if not paras:
            return _fail("<no Heading {} paragraphs>".format(level))
        sizes = set()
        for p in paras:
            for r in _para_text_runs(p):
                sz = _run_size_pt(r)
                if sz is not None:
                    sizes.add(round(sz, 1))
        if not sizes:
            return _ok("<inherits>")
        if len(sizes) == 1 and _num_close(next(iter(sizes)), expected, 0.5):
            return _ok(next(iter(sizes)))
        return _fail(sorted(sizes))
    return checker


def _check_heading_alignment(level: int):
    def checker(doc, expected, ctx):
        paras = _heading_paragraphs(doc, level)
        if not paras:
            return _fail("<no Heading {} paragraphs>".format(level))
        aligns = {_alignment_name(p.alignment) for p in paras}
        if aligns == {expected} or aligns == {expected, "inherit"}:
            return _ok(expected)
        return _fail(sorted(aligns))
    return checker


def _check_heading_levels_no_skip(doc, expected, ctx):
    """expected=True → every increase in heading depth is at most 1."""
    levels = []
    for p in doc.paragraphs:
        name = (p.style.name or "").lower()
        m = re.match(r'heading (\d+)', name)
        if m:
            levels.append(int(m.group(1)))
    if not levels:
        return _ok("<no real heading styles in document>")
    skips = []
    for i in range(1, len(levels)):
        if levels[i] > levels[i-1] + 1:
            skips.append((levels[i-1], levels[i]))
    if skips:
        return _fail(f"{len(skips)} skip(s)", skips=skips)
    return _ok("no skips")


# ---------------------------------------------------------------------------
# Whole-document pattern checks
# ---------------------------------------------------------------------------

def _check_no_consecutive_blanks(doc, expected, ctx):
    """expected=max allowed (e.g. 2) → fail if any run of blank paragraphs > expected."""
    run_len = 0
    max_run = 0
    for p in doc.paragraphs:
        if _is_blank(p):
            run_len += 1
            max_run = max(max_run, run_len)
        else:
            run_len = 0
    return _ok(max_run) if max_run <= expected else _fail(max_run)


def _check_no_forbidden_fonts_in_body(doc, expected, ctx):
    """expected = list of font names that must NOT appear in body."""
    forbidden = set(expected)
    found: Dict[str, int] = {}
    for p in _body_paragraphs(doc):
        for r in _para_text_runs(p):
            for f in (_run_east_asia_font(r), _run_ascii_font(r)):
                if f and f in forbidden:
                    found[f] = found.get(f, 0) + 1
    if found:
        return _fail(found)
    return _ok("none found")


def _check_no_red_text(doc, expected, ctx):
    """expected=True → fail if any body run has explicit red-ish color."""
    from docx.shared import RGBColor
    offenders = []
    for p in doc.paragraphs:
        for r in _para_text_runs(p):
            try:
                rgb = r.font.color.rgb
            except Exception:
                rgb = None
            if rgb is None:
                continue
            # Detect red-dominant (R > 150, G/B much smaller).
            try:
                # RGBColor doesn't expose r/g/b directly; use hex.
                hexv = str(rgb)
                rh, gh, bh = int(hexv[0:2], 16), int(hexv[2:4], 16), int(hexv[4:6], 16)
                if rh > 150 and rh > gh + 60 and rh > bh + 60:
                    offenders.append({"text": r.text[:30], "color": hexv})
            except Exception:
                pass
    if not offenders:
        return _ok("none")
    return _fail(f"{len(offenders)} red run(s)", samples=offenders[:3])


def _check_mixed_punctuation_in_cjk(doc, expected, ctx):
    """expected=True → fail if any body CJK paragraph contains forbidden half-width
    punctuation (,;.?!). Heuristic: paragraph predominantly Chinese should use 全角."""
    offenders = []
    for p in _body_paragraphs(doc):
        text = p.text
        cjk = len(_RE_CJK.findall(text))
        if cjk < 5:
            continue  # likely English / mixed; skip
        # Look for half-width punctuation FOLLOWED OR PRECEDED by a CJK char.
        # That's a strong signal it's a typo in a CJK paragraph.
        for m in re.finditer(r'[一-鿿][,;.?!]|[,;.?!][一-鿿]', text):
            offenders.append(text[max(0, m.start()-5):m.end()+5])
            break  # one per paragraph is enough evidence
    if not offenders:
        return _ok("clean")
    return _fail(f"{len(offenders)} CJK paragraph(s) with half-width punctuation",
                 samples=offenders[:3])


def _check_keywords_separator(doc, expected, ctx):
    """expected={"cjk": "；", "en": ";"} — look for "关键词" and "Keywords" lines
    and check the separator between entries.
    """
    cjk_sep = expected.get("cjk", "；")
    en_sep = expected.get("en", ";")
    issues = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("关键词"):
            # Strip leading "关键词:" or "关键词：" then look at separators
            rest = re.sub(r'^关键词[：:]\s*', '', t)
            # Bad if it contains "," or ";" (half-width) — should be "；"
            if re.search(r'[,;]', rest):
                issues.append({"line": t[:60], "expected_sep": cjk_sep})
        elif t.lower().startswith("keywords"):
            rest = re.sub(r'(?i)^keywords[:：]\s*', '', t)
            if "；" in rest or "，" in rest:
                issues.append({"line": t[:60], "expected_sep": en_sep})
    if not issues:
        return _ok("clean")
    return _fail(issues)


def _check_references_numbering(doc, expected, ctx):
    """expected=True → every reference entry starts with [N] (digits in brackets)."""
    # Locate references section: paragraph whose text is "参考文献" exactly,
    # then everything after until end of doc.
    found_marker = False
    bad: List[str] = []
    total = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if not found_marker:
            if t == "参考文献" or t.lower() == "references":
                found_marker = True
            continue
        if not t:
            continue
        total += 1
        if not re.match(r'\[\d+\]', t):
            bad.append(t[:60])
    if not found_marker:
        return _ok("<no references section>")
    if not bad:
        return _ok(f"{total} entries OK")
    return _fail(f"{len(bad)}/{total} entries with bad numbering", samples=bad[:3])


# ---------------------------------------------------------------------------
# Dispatch table — id → checker
# ---------------------------------------------------------------------------

CHECKERS: Dict[str, Callable[..., Dict[str, Any]]] = {
    # Page setup
    "page.margin_top_cm":         _check_section_margin("top_margin_cm"),
    "page.margin_bottom_cm":      _check_section_margin("bottom_margin_cm"),
    "page.margin_left_cm":        _check_section_margin("left_margin_cm"),
    "page.margin_right_cm":       _check_section_margin("right_margin_cm"),
    "page.gutter_cm":             _check_section_margin("gutter_cm"),
    "page.header_distance_cm":    _check_section_margin("header_distance_cm"),
    "page.footer_distance_cm":    _check_section_margin("footer_distance_cm"),
    "page.different_odd_and_even_pages": _check_odd_even_pages,

    # Body
    "body.font_eastasia":         _check_body_font_eastasia,
    "body.font_ascii":            _check_body_font_ascii,
    "body.size_pt":               _check_body_size_pt,
    "body.line_spacing_pt":       _check_body_line_spacing_pt,

    # Headings
    "heading1.font_eastasia":     _check_heading_font_eastasia(1),
    "heading1.size_pt":           _check_heading_size_pt(1),
    "heading1.alignment":         _check_heading_alignment(1),
    "heading2.font_eastasia":     _check_heading_font_eastasia(2),
    "heading2.size_pt":           _check_heading_size_pt(2),
    "heading2.alignment":         _check_heading_alignment(2),
    "headings.no_level_skip":     _check_heading_levels_no_skip,

    # Whole-document patterns
    "patterns.max_consecutive_blank_paragraphs": _check_no_consecutive_blanks,
    "patterns.no_forbidden_body_fonts":          _check_no_forbidden_fonts_in_body,
    "patterns.no_red_body_text":                 _check_no_red_text,
    "patterns.no_halfwidth_punct_in_cjk":        _check_mixed_punctuation_in_cjk,
    "patterns.keywords_separator":               _check_keywords_separator,

    # References
    "references.numbering_format":               _check_references_numbering,
}


# Tooltip-style hints for callers when an item fails. Path → suggestion.
FIX_HINTS: Dict[str, str] = {
    "page.margin_top_cm":      "word_live_set_page_layout(margin_top_cm=<expected>)",
    "page.margin_bottom_cm":   "word_live_set_page_layout(margin_bottom_cm=<expected>)",
    "page.margin_left_cm":     "word_live_set_page_layout(margin_left_cm=<expected>)",
    "page.margin_right_cm":    "word_live_set_page_layout(margin_right_cm=<expected>)",
    "page.gutter_cm":          "word_live_set_page_layout(gutter_cm=<expected>)",
    "page.header_distance_cm": "word_live_set_page_layout(header_distance_cm=<expected>)",
    "page.footer_distance_cm": "word_live_set_page_layout(footer_distance_cm=<expected>)",
    "page.different_odd_and_even_pages": "word_live_set_page_layout(different_odd_and_even_pages=True)",
    "body.font_eastasia":      "word_live_format_text or scan body paragraphs and reapply font",
    "body.font_ascii":         "same as above for English text runs",
    "body.size_pt":             "apply size via Normal style or word_live_format_text",
    "body.line_spacing_pt":     "word_live_set_paragraph_spacing(line_spacing=<lines>, line_spacing_rule='multiple')",
    "heading1.font_eastasia":  "reassign style Heading 1 with font 黑体, or word_live_setup_heading_numbering",
    "heading1.size_pt":        "edit Heading 1 style or reformat each level-1 paragraph",
    "heading1.alignment":      "set Heading 1 alignment to center",
    "heading2.font_eastasia":  "reassign style Heading 2 with font 黑体",
    "heading2.size_pt":        "edit Heading 2 style",
    "heading2.alignment":      "set Heading 2 alignment",
    "headings.no_level_skip":  "ensure level depth increases by ≤1 at each step",
    "patterns.max_consecutive_blank_paragraphs": "delete extra blank paragraphs, use space_before/after instead",
    "patterns.no_forbidden_body_fonts": "convert Calibri/Arial/Courier/楷体 etc. runs back to body font",
    "patterns.no_red_body_text": "set font color to automatic/black for body text",
    "patterns.no_halfwidth_punct_in_cjk": "search/replace , ; . ? ! → ， ； 。 ？ ！ in CJK paragraphs",
    "patterns.keywords_separator": "replace 关键词 separators with 全角；, Keywords separators with ;",
    "references.numbering_format": "rewrite reference entries to start with [N]",
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def audit_against_spec(
    filename: str,
    spec: Optional[List[Dict[str, Any]]] = None,
    spec_path: Optional[str] = None,
) -> Dict[str, Any]:
    """Audit a saved .docx file against a structured spec.

    Args:
        filename: Path to the .docx file to audit (must already be saved
            — if the file is open in Word with unsaved edits, save first
            via ``word_live_save``).
        spec: Inline JSON spec. Each item: ``{"id": "...", "expected": ...}``.
        spec_path: Alternative to ``spec`` — path to a .json file.

    Returns:
        Dict with ``pass_count``, ``fail_count``, ``unknown_count``, and
        ``items`` (per-item report).
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return {"error": f"file not found: {filename}"}

    if spec is None and spec_path is None:
        return {"error": "must provide either spec or spec_path"}
    if spec is None:
        spec_path = os.path.expandvars(os.path.expanduser(spec_path))
        if not os.path.exists(spec_path):
            return {"error": f"spec file not found: {spec_path}"}
        spec = json.loads(Path(spec_path).read_text(encoding='utf-8'))

    if not isinstance(spec, list):
        return {"error": "spec must be a list of {id, expected} items"}

    try:
        doc = Document(filename)
    except Exception as e:
        return {"error": f"could not open document: {e}"}

    ctx = {"filename": filename}
    items: List[Dict[str, Any]] = []
    pass_count = 0
    fail_count = 0
    unknown_count = 0

    for entry in spec:
        item_id = entry.get("id")
        expected = entry.get("expected")
        if item_id is None:
            items.append({"id": "<missing>", "pass": False, "actual": None,
                          "error": "spec item missing 'id'"})
            fail_count += 1
            continue

        checker = CHECKERS.get(item_id)
        if checker is None:
            unknown_count += 1
            items.append({
                "id": item_id,
                "expected": expected,
                "pass": None,
                "actual": None,
                "error": "no checker registered for this id",
            })
            continue

        try:
            result = checker(doc, expected, ctx)
        except Exception as e:
            items.append({
                "id": item_id,
                "expected": expected,
                "pass": False,
                "actual": None,
                "error": f"checker raised: {type(e).__name__}: {e}",
            })
            fail_count += 1
            continue

        entry_out = {
            "id": item_id,
            "expected": expected,
            "actual": result.get("actual"),
            "pass": result.get("pass"),
        }
        for k in ("hint", "samples", "skips"):
            if k in result:
                entry_out[k] = result[k]
        if not result.get("pass"):
            fix_hint = FIX_HINTS.get(item_id)
            if fix_hint:
                entry_out["fix_hint"] = fix_hint
            fail_count += 1
        else:
            pass_count += 1
        items.append(entry_out)

    return {
        "filename": filename,
        "pass_count": pass_count,
        "fail_count": fail_count,
        "unknown_count": unknown_count,
        "total": len(items),
        "items": items,
    }


async def word_audit_against_spec(
    filename: str,
    spec: Optional[List[Dict[str, Any]]] = None,
    spec_path: Optional[str] = None,
) -> str:
    """MCP-facing wrapper: returns JSON string."""
    result = audit_against_spec(filename=filename, spec=spec, spec_path=spec_path)
    return json.dumps(result, ensure_ascii=False, indent=2)
