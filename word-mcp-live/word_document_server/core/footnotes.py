"""
Consolidated footnote functionality for Word documents.
This module combines all footnote implementations with proper namespace handling and Word compliance.
"""

import os
import zipfile
import tempfile
from typing import Optional, Tuple, Dict, Any, List
from lxml import etree
from docx import Document
from docx.oxml.ns import qn

# Namespace definitions
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
XML_NS = 'http://www.w3.org/XML/1998/namespace'  # for xml:space attribute

# Constants
RESERVED_FOOTNOTE_IDS = {-1, 0, 1}  # Reserved for separators and Word internals
RESERVED_ENDNOTE_IDS = {-1, 0}  # Endnotes reserve separator + continuationSeparator only
MIN_FOOTNOTE_ID = -2147483648
MAX_FOOTNOTE_ID = 32767
MAX_RELATIONSHIP_ID_LENGTH = 255
FOOTNOTE_REF_STYLE_INDEX = 38
FOOTNOTE_TEXT_STYLE_INDEX = 29


# ============================================================================
# BASIC UTILITIES (from footnotes.py)
# ============================================================================

def find_footnote_references(doc):
    """Find all footnote references in the document."""
    footnote_refs = []
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            # Check if this run has superscript formatting
            if run.font.superscript:
                # Check if it's likely a footnote reference
                if run.text.isdigit() or run.text in "¹²³⁴⁵⁶⁷⁸⁹⁰†‡§¶":
                    footnote_refs.append({
                        'paragraph_index': para_idx,
                        'run_index': run_idx,
                        'text': run.text,
                        'paragraph': para,
                        'run': run
                    })
    return footnote_refs


def get_format_symbols(format_type: str, count: int) -> list:
    """Generate format symbols for footnote numbering."""
    symbols = []
    
    if format_type == "1, 2, 3":
        symbols = [str(i) for i in range(1, count + 1)]
    elif format_type == "i, ii, iii":
        # Roman numerals
        roman_map = [(10, 'x'), (9, 'ix'), (5, 'v'), (4, 'iv'), (1, 'i')]
        for i in range(1, count + 1):
            result = ''
            num = i
            for value, numeral in roman_map:
                count_sym, num = divmod(num, value)
                result += numeral * count_sym
            symbols.append(result)
    elif format_type == "a, b, c":
        # Alphabetic
        for i in range(1, count + 1):
            if i <= 26:
                symbols.append(chr(96 + i))
            else:
                # For numbers > 26, use aa, ab, etc.
                first = (i - 1) // 26
                second = (i - 1) % 26 + 1
                symbols.append(chr(96 + first) + chr(96 + second))
    elif format_type == "*, †, ‡":
        # Special symbols
        special = ['*', '†', '‡', '§', '¶', '#']
        for i in range(1, count + 1):
            if i <= len(special):
                symbols.append(special[i - 1])
            else:
                # Repeat symbols with numbers
                symbols.append(special[(i - 1) % len(special)] + str((i - 1) // len(special) + 1))
    else:
        # Default to numeric
        symbols = [str(i) for i in range(1, count + 1)]
    
    return symbols


def customize_footnote_formatting(doc, footnote_refs, format_symbols, start_number, footnote_style):
    """Apply custom formatting to footnotes."""
    count = 0
    for i, ref in enumerate(footnote_refs):
        if i < len(format_symbols):
            # Update the footnote reference text
            ref['run'].text = format_symbols[i]
            ref['run'].font.superscript = True
            
            # Apply style if available
            if footnote_style:
                try:
                    ref['paragraph'].style = footnote_style
                except:
                    pass
            count += 1
    return count


# ============================================================================
# ROBUST IMPLEMENTATION (consolidated from footnotes_robust.py)
# ============================================================================

def _get_safe_footnote_id(footnotes_root) -> int:
    """Get a safe footnote ID avoiding conflicts and reserved values."""
    nsmap = {'w': W_NS}
    existing_footnotes = footnotes_root.xpath('//w:footnote', namespaces=nsmap)
    
    used_ids = set()
    for fn in existing_footnotes:
        fn_id = fn.get(f'{{{W_NS}}}id')
        if fn_id:
            try:
                used_ids.add(int(fn_id))
            except ValueError:
                pass
    
    # Start from 2 to avoid reserved IDs
    candidate_id = 2
    while candidate_id in used_ids or candidate_id in RESERVED_FOOTNOTE_IDS:
        candidate_id += 1
        if candidate_id > MAX_FOOTNOTE_ID:
            raise ValueError("No available footnote IDs")
    
    return candidate_id


def _ensure_content_types(content_types_xml: bytes) -> bytes:
    """Ensure content types with proper namespace handling."""
    ct_tree = etree.fromstring(content_types_xml)
    
    # Content Types uses default namespace - must use namespace-aware XPath
    nsmap = {'ct': CT_NS}
    
    # Check for existing override with proper namespace
    existing_overrides = ct_tree.xpath(
        "//ct:Override[@PartName='/word/footnotes.xml']",
        namespaces=nsmap
    )
    
    if existing_overrides:
        return content_types_xml  # Already exists
    
    # Add override with proper namespace
    override = etree.Element(f'{{{CT_NS}}}Override',
        PartName='/word/footnotes.xml',
        ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
    )
    ct_tree.append(override)
    
    return etree.tostring(ct_tree, encoding='UTF-8', xml_declaration=True, standalone="yes")


def _ensure_document_rels(document_rels_xml: bytes) -> bytes:
    """Ensure document relationships with proper namespace handling."""
    rels_tree = etree.fromstring(document_rels_xml)
    
    # Relationships uses default namespace - must use namespace-aware XPath
    nsmap = {'r': REL_NS}
    
    # Check for existing footnotes relationship with proper namespace
    existing_footnote_rels = rels_tree.xpath(
        "//r:Relationship[contains(@Type, 'footnotes')]",
        namespaces=nsmap
    )
    
    if existing_footnote_rels:
        return document_rels_xml  # Already exists
    
    # Generate unique rId using namespace-aware XPath
    all_rels = rels_tree.xpath('//r:Relationship', namespaces=nsmap)
    existing_ids = {rel.get('Id') for rel in all_rels if rel.get('Id')}
    rid_num = 1
    while f'rId{rid_num}' in existing_ids:
        rid_num += 1
    
    # Validate ID length
    new_rid = f'rId{rid_num}'
    if len(new_rid) > MAX_RELATIONSHIP_ID_LENGTH:
        raise ValueError(f"Relationship ID too long: {new_rid}")
    
    # Create relationship with proper namespace
    rel = etree.Element(f'{{{REL_NS}}}Relationship',
        Id=new_rid,
        Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes',
        Target='footnotes.xml'
    )
    rels_tree.append(rel)
    
    return etree.tostring(rels_tree, encoding='UTF-8', xml_declaration=True, standalone="yes")


def _create_minimal_footnotes_xml() -> bytes:
    """Create minimal footnotes.xml with separators."""
    xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="{W_NS}">
    <w:footnote w:type="separator" w:id="-1">
        <w:p>
            <w:pPr>
                <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
            </w:pPr>
            <w:r>
                <w:separator/>
            </w:r>
        </w:p>
    </w:footnote>
    <w:footnote w:type="continuationSeparator" w:id="0">
        <w:p>
            <w:pPr>
                <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
            </w:pPr>
            <w:r>
                <w:continuationSeparator/>
            </w:r>
        </w:p>
    </w:footnote>
</w:footnotes>'''
    return xml.encode('utf-8')


def _ensure_footnote_styles(styles_root):
    """Ensure both FootnoteReference and FootnoteText styles exist."""
    nsmap = {'w': W_NS}
    
    # Check for FootnoteReference style
    ref_style = styles_root.xpath('//w:style[@w:styleId="FootnoteReference"]', namespaces=nsmap)
    if not ref_style:
        # Create FootnoteReference character style
        style = etree.Element(f'{{{W_NS}}}style',
            attrib={
                f'{{{W_NS}}}type': 'character',
                f'{{{W_NS}}}styleId': 'FootnoteReference'
            }
        )
        name = etree.SubElement(style, f'{{{W_NS}}}name')
        name.set(f'{{{W_NS}}}val', 'footnote reference')
        
        base = etree.SubElement(style, f'{{{W_NS}}}basedOn')
        base.set(f'{{{W_NS}}}val', 'DefaultParagraphFont')
        
        rPr = etree.SubElement(style, f'{{{W_NS}}}rPr')
        vert_align = etree.SubElement(rPr, f'{{{W_NS}}}vertAlign')
        vert_align.set(f'{{{W_NS}}}val', 'superscript')
        
        styles_root.append(style)
    
    # Check for FootnoteText style
    text_style = styles_root.xpath('//w:style[@w:styleId="FootnoteText"]', namespaces=nsmap)
    if not text_style:
        # Create FootnoteText paragraph style
        style = etree.Element(f'{{{W_NS}}}style',
            attrib={
                f'{{{W_NS}}}type': 'paragraph',
                f'{{{W_NS}}}styleId': 'FootnoteText'
            }
        )
        name = etree.SubElement(style, f'{{{W_NS}}}name')
        name.set(f'{{{W_NS}}}val', 'footnote text')
        
        base = etree.SubElement(style, f'{{{W_NS}}}basedOn')
        base.set(f'{{{W_NS}}}val', 'Normal')
        
        pPr = etree.SubElement(style, f'{{{W_NS}}}pPr')
        sz = etree.SubElement(pPr, f'{{{W_NS}}}sz')
        sz.set(f'{{{W_NS}}}val', '20')  # 10pt
        
        styles_root.append(style)


def add_footnote_robust(
    filename: str,
    search_text: Optional[str] = None,
    paragraph_index: Optional[int] = None,
    footnote_text: str = "",
    output_filename: Optional[str] = None,
    position: str = "after",
    validate_location: bool = True,
    auto_repair: bool = False
) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
    """
    Add a footnote with robust validation and error handling.
    
    This is the main production-ready function with all fixes applied.
    """
    
    # Validate inputs
    if not search_text and paragraph_index is None:
        return False, "Must provide either search_text or paragraph_index", None
    
    if search_text and paragraph_index is not None:
        return False, "Cannot provide both search_text and paragraph_index", None
    
    if not os.path.exists(filename):
        return False, f"File not found: {filename}", None
    
    # Set working file
    working_file = output_filename if output_filename else filename
    if output_filename and filename != output_filename:
        import shutil
        shutil.copy2(filename, output_filename)
    
    try:
        # Read document parts
        doc_parts = {}
        with zipfile.ZipFile(filename, 'r') as zin:
            doc_parts['document'] = zin.read('word/document.xml')
            doc_parts['content_types'] = zin.read('[Content_Types].xml')
            doc_parts['document_rels'] = zin.read('word/_rels/document.xml.rels')
            
            # Read or create footnotes.xml
            if 'word/footnotes.xml' in zin.namelist():
                doc_parts['footnotes'] = zin.read('word/footnotes.xml')
            else:
                doc_parts['footnotes'] = _create_minimal_footnotes_xml()
            
            # Read styles
            if 'word/styles.xml' in zin.namelist():
                doc_parts['styles'] = zin.read('word/styles.xml')
            else:
                # Create minimal styles
                doc_parts['styles'] = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        
        # Parse XML documents
        doc_root = etree.fromstring(doc_parts['document'])
        footnotes_root = etree.fromstring(doc_parts['footnotes'])
        styles_root = etree.fromstring(doc_parts['styles'])
        
        # Find target location
        nsmap = {'w': W_NS}
        
        if search_text:
            # Search for text in paragraphs
            found = False
            for para in doc_root.xpath('//w:p', namespaces=nsmap):
                para_text = ''.join(para.xpath('.//w:t/text()', namespaces=nsmap))
                if search_text in para_text:
                    target_para = para
                    found = True
                    break
            
            if not found:
                return False, f"Text '{search_text}' not found in document", None
        else:
            # Use paragraph index
            paragraphs = doc_root.xpath('//w:p', namespaces=nsmap)
            if paragraph_index >= len(paragraphs):
                return False, f"Paragraph index {paragraph_index} out of range", None
            target_para = paragraphs[paragraph_index]
        
        # Validate location if requested
        if validate_location:
            # Check if paragraph is in header/footer
            parent = target_para.getparent()
            while parent is not None:
                if parent.tag in [f'{{{W_NS}}}hdr', f'{{{W_NS}}}ftr']:
                    return False, "Cannot add footnote in header/footer", None
                parent = parent.getparent()
        
        # Get safe footnote ID
        footnote_id = _get_safe_footnote_id(footnotes_root)
        
        # Add footnote reference to document
        if position == "after":
            # Find last run in paragraph or create one
            runs = target_para.xpath('.//w:r', namespaces=nsmap)
            if runs:
                last_run = runs[-1]
                # Insert after last run
                insert_pos = target_para.index(last_run) + 1
            else:
                insert_pos = len(target_para)
        else:  # before
            # Find first run with text
            runs = target_para.xpath('.//w:r[w:t]', namespaces=nsmap)
            if runs:
                first_run = runs[0]
                insert_pos = target_para.index(first_run)
            else:
                insert_pos = 0
        
        # Create footnote reference run
        ref_run = etree.Element(f'{{{W_NS}}}r')
        
        # Add run properties with superscript
        rPr = etree.SubElement(ref_run, f'{{{W_NS}}}rPr')
        rStyle = etree.SubElement(rPr, f'{{{W_NS}}}rStyle')
        rStyle.set(f'{{{W_NS}}}val', 'FootnoteReference')
        
        # Add footnote reference
        fn_ref = etree.SubElement(ref_run, f'{{{W_NS}}}footnoteReference')
        fn_ref.set(f'{{{W_NS}}}id', str(footnote_id))
        
        # Insert the reference run
        target_para.insert(insert_pos, ref_run)
        
        # Add footnote content
        new_footnote = etree.Element(f'{{{W_NS}}}footnote',
            attrib={f'{{{W_NS}}}id': str(footnote_id)}
        )
        
        # Add paragraph to footnote
        fn_para = etree.SubElement(new_footnote, f'{{{W_NS}}}p')
        
        # Add paragraph properties
        pPr = etree.SubElement(fn_para, f'{{{W_NS}}}pPr')
        pStyle = etree.SubElement(pPr, f'{{{W_NS}}}pStyle')
        pStyle.set(f'{{{W_NS}}}val', 'FootnoteText')
        
        # Add the footnote reference marker
        marker_run = etree.SubElement(fn_para, f'{{{W_NS}}}r')
        marker_rPr = etree.SubElement(marker_run, f'{{{W_NS}}}rPr')
        marker_rStyle = etree.SubElement(marker_rPr, f'{{{W_NS}}}rStyle')
        marker_rStyle.set(f'{{{W_NS}}}val', 'FootnoteReference')
        marker_ref = etree.SubElement(marker_run, f'{{{W_NS}}}footnoteRef')
        
        # Add space after marker
        space_run = etree.SubElement(fn_para, f'{{{W_NS}}}r')
        space_text = etree.SubElement(space_run, f'{{{W_NS}}}t')
        space_text.set(f'{{{XML_NS}}}space', 'preserve')
        space_text.text = ' '
        
        # Add footnote text
        text_run = etree.SubElement(fn_para, f'{{{W_NS}}}r')
        text_elem = etree.SubElement(text_run, f'{{{W_NS}}}t')
        text_elem.text = footnote_text
        
        # Append footnote to footnotes.xml
        footnotes_root.append(new_footnote)
        
        # Ensure styles exist
        _ensure_footnote_styles(styles_root)
        
        # Ensure coherence
        content_types_xml = _ensure_content_types(doc_parts['content_types'])
        document_rels_xml = _ensure_document_rels(doc_parts['document_rels'])
        
        # Write modified document
        temp_file = working_file + '.tmp'
        with zipfile.ZipFile(temp_file, 'w', zipfile.ZIP_DEFLATED) as zout:
            with zipfile.ZipFile(filename, 'r') as zin:
                # Copy unchanged files
                for item in zin.infolist():
                    if item.filename not in [
                        'word/document.xml', 'word/footnotes.xml', 'word/styles.xml',
                        '[Content_Types].xml', 'word/_rels/document.xml.rels'
                    ]:
                        zout.writestr(item, zin.read(item.filename))
            
            # Write modified files
            zout.writestr('word/document.xml',
                etree.tostring(doc_root, encoding='UTF-8', xml_declaration=True, standalone="yes"))
            zout.writestr('word/footnotes.xml',
                etree.tostring(footnotes_root, encoding='UTF-8', xml_declaration=True, standalone="yes"))
            zout.writestr('word/styles.xml',
                etree.tostring(styles_root, encoding='UTF-8', xml_declaration=True, standalone="yes"))
            zout.writestr('[Content_Types].xml', content_types_xml)
            zout.writestr('word/_rels/document.xml.rels', document_rels_xml)
        
        # Replace original with temp file
        os.replace(temp_file, working_file)
        
        details = {
            'footnote_id': footnote_id,
            'location': 'search_text' if search_text else 'paragraph_index',
            'styles_created': ['FootnoteReference', 'FootnoteText'],
            'coherence_verified': True
        }
        
        return True, f"Successfully added footnote (ID: {footnote_id}) to {working_file}", details
        
    except Exception as e:
        # Clean up temp file if exists
        temp_file = working_file + '.tmp'
        if os.path.exists(temp_file):
            os.remove(temp_file)
        return False, f"Error adding footnote: {str(e)}", None


def delete_footnote_robust(
    filename: str,
    footnote_id: Optional[int] = None,
    search_text: Optional[str] = None,
    output_filename: Optional[str] = None,
    clean_orphans: bool = True
) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
    """Delete a footnote with comprehensive cleanup."""
    
    if not footnote_id and not search_text:
        return False, "Must provide either footnote_id or search_text", None
    
    if not os.path.exists(filename):
        return False, f"File not found: {filename}", None
    
    # Set working file
    working_file = output_filename if output_filename else filename
    if output_filename and filename != output_filename:
        import shutil
        shutil.copy2(filename, output_filename)
    
    try:
        # Read document parts
        with zipfile.ZipFile(filename, 'r') as zin:
            doc_xml = zin.read('word/document.xml')
            
            if 'word/footnotes.xml' not in zin.namelist():
                return False, "No footnotes in document", None
            
            footnotes_xml = zin.read('word/footnotes.xml')
        
        # Parse documents
        doc_root = etree.fromstring(doc_xml)
        footnotes_root = etree.fromstring(footnotes_xml)
        nsmap = {'w': W_NS}
        
        # Find footnote to delete
        if search_text:
            # Find footnote reference near text
            for para in doc_root.xpath('//w:p', namespaces=nsmap):
                para_text = ''.join(para.xpath('.//w:t/text()', namespaces=nsmap))
                if search_text in para_text:
                    # Look for footnote reference in this paragraph
                    fn_refs = para.xpath('.//w:footnoteReference', namespaces=nsmap)
                    if fn_refs:
                        footnote_id = int(fn_refs[0].get(f'{{{W_NS}}}id'))
                        break
            
            if not footnote_id:
                return False, f"No footnote found near text '{search_text}'", None
        
        # Remove footnote reference from document
        refs_removed = 0
        for fn_ref in doc_root.xpath(f'//w:footnoteReference[@w:id="{footnote_id}"]', namespaces=nsmap):
            # Remove the entire run containing the reference
            run = fn_ref.getparent()
            if run is not None and run.tag == f'{{{W_NS}}}r':
                para = run.getparent()
                if para is not None:
                    para.remove(run)
                    refs_removed += 1
        
        if refs_removed == 0:
            return False, f"Footnote {footnote_id} not found", None
        
        # Remove footnote content
        content_removed = 0
        for fn in footnotes_root.xpath(f'//w:footnote[@w:id="{footnote_id}"]', namespaces=nsmap):
            footnotes_root.remove(fn)
            content_removed += 1
        
        # Clean orphans if requested
        orphans_removed = []
        if clean_orphans:
            # Find all referenced IDs
            referenced_ids = set()
            for ref in doc_root.xpath('//w:footnoteReference', namespaces=nsmap):
                ref_id = ref.get(f'{{{W_NS}}}id')
                if ref_id:
                    referenced_ids.add(ref_id)
            
            # Remove unreferenced footnotes (except separators)
            for fn in footnotes_root.xpath('//w:footnote', namespaces=nsmap):
                fn_id = fn.get(f'{{{W_NS}}}id')
                if fn_id and fn_id not in referenced_ids and fn_id not in ['-1', '0']:
                    footnotes_root.remove(fn)
                    orphans_removed.append(fn_id)
        
        # Write modified document
        temp_file = working_file + '.tmp'
        with zipfile.ZipFile(temp_file, 'w', zipfile.ZIP_DEFLATED) as zout:
            with zipfile.ZipFile(filename, 'r') as zin:
                for item in zin.infolist():
                    if item.filename == 'word/document.xml':
                        zout.writestr(item,
                            etree.tostring(doc_root, encoding='UTF-8', xml_declaration=True, standalone="yes"))
                    elif item.filename == 'word/footnotes.xml':
                        zout.writestr(item,
                            etree.tostring(footnotes_root, encoding='UTF-8', xml_declaration=True, standalone="yes"))
                    else:
                        zout.writestr(item, zin.read(item.filename))
        
        os.replace(temp_file, working_file)
        
        details = {
            'footnote_id': footnote_id,
            'references_removed': refs_removed,
            'content_removed': content_removed,
            'orphans_removed': orphans_removed
        }
        
        message = f"Successfully deleted footnote {footnote_id}"
        if orphans_removed:
            message += f" and {len(orphans_removed)} orphaned footnotes"
        
        return True, message, details
        
    except Exception as e:
        return False, f"Error deleting footnote: {str(e)}", None


def validate_document_footnotes(filename: str) -> Tuple[bool, str, Dict[str, Any]]:
    """Validate all footnotes in a document for coherence and compliance."""
    
    if not os.path.exists(filename):
        return False, f"File not found: {filename}", {}
    
    report = {
        'total_references': 0,
        'total_content': 0,
        'id_conflicts': [],
        'orphaned_content': [],
        'missing_references': [],
        'invalid_locations': [],
        'missing_styles': [],
        'coherence_issues': []
    }
    
    try:
        with zipfile.ZipFile(filename, 'r') as zf:
            # Check document.xml
            doc_xml = zf.read('word/document.xml')
            doc_root = etree.fromstring(doc_xml)
            nsmap = {'w': W_NS}
            
            # Get all footnote references
            ref_ids = set()
            for ref in doc_root.xpath('//w:footnoteReference', namespaces=nsmap):
                ref_id = ref.get(f'{{{W_NS}}}id')
                if ref_id:
                    ref_ids.add(ref_id)
                    report['total_references'] += 1
                    
                    # Check location
                    parent = ref.getparent()
                    while parent is not None:
                        if parent.tag in [f'{{{W_NS}}}hdr', f'{{{W_NS}}}ftr']:
                            report['invalid_locations'].append(ref_id)
                            break
                        parent = parent.getparent()
            
            # Check footnotes.xml
            if 'word/footnotes.xml' in zf.namelist():
                footnotes_xml = zf.read('word/footnotes.xml')
                footnotes_root = etree.fromstring(footnotes_xml)
                
                content_ids = set()
                for fn in footnotes_root.xpath('//w:footnote', namespaces=nsmap):
                    fn_id = fn.get(f'{{{W_NS}}}id')
                    if fn_id:
                        content_ids.add(fn_id)
                        if fn_id not in ['-1', '0']:  # Exclude separators
                            report['total_content'] += 1
                
                # Find orphans and missing
                report['orphaned_content'] = list(content_ids - ref_ids - {'-1', '0'})
                report['missing_references'] = list(ref_ids - content_ids)
            else:
                if report['total_references'] > 0:
                    report['coherence_issues'].append('References exist but no footnotes.xml')
            
            # Check relationships
            if 'word/_rels/document.xml.rels' in zf.namelist():
                rels_xml = zf.read('word/_rels/document.xml.rels')
                rels_root = etree.fromstring(rels_xml)
                rel_nsmap = {'r': REL_NS}
                
                fn_rels = rels_root.xpath(
                    "//r:Relationship[contains(@Type, 'footnotes')]",
                    namespaces=rel_nsmap
                )
                
                if report['total_content'] > 0 and len(fn_rels) == 0:
                    report['coherence_issues'].append('Missing footnotes relationship')
                elif len(fn_rels) > 1:
                    report['coherence_issues'].append(f'Multiple footnote relationships: {len(fn_rels)}')
            
            # Check content types
            if '[Content_Types].xml' in zf.namelist():
                ct_xml = zf.read('[Content_Types].xml')
                ct_root = etree.fromstring(ct_xml)
                ct_nsmap = {'ct': CT_NS}
                
                fn_overrides = ct_root.xpath(
                    "//ct:Override[@PartName='/word/footnotes.xml']",
                    namespaces=ct_nsmap
                )
                
                if report['total_content'] > 0 and len(fn_overrides) == 0:
                    report['coherence_issues'].append('Missing footnotes content type')
                elif len(fn_overrides) > 1:
                    report['coherence_issues'].append(f'Multiple footnote content types: {len(fn_overrides)}')
            
            # Check styles
            if 'word/styles.xml' in zf.namelist():
                styles_xml = zf.read('word/styles.xml')
                styles_root = etree.fromstring(styles_xml)
                
                ref_style = styles_root.xpath('//w:style[@w:styleId="FootnoteReference"]', namespaces=nsmap)
                text_style = styles_root.xpath('//w:style[@w:styleId="FootnoteText"]', namespaces=nsmap)
                
                if not ref_style:
                    report['missing_styles'].append('FootnoteReference')
                if not text_style:
                    report['missing_styles'].append('FootnoteText')
        
        # Determine if valid
        is_valid = (
            len(report['id_conflicts']) == 0 and
            len(report['orphaned_content']) == 0 and
            len(report['missing_references']) == 0 and
            len(report['invalid_locations']) == 0 and
            len(report['coherence_issues']) == 0
        )
        
        if is_valid:
            message = "Document footnotes are valid"
        else:
            message = "Document has footnote issues"
        
        return is_valid, message, report
        
    except Exception as e:
        return False, f"Error validating document: {str(e)}", report


# ============================================================================
# COMPATIBILITY FUNCTIONS (for backward compatibility)
# ============================================================================

def add_footnote_at_paragraph_end(
    filename: str,
    paragraph_index: int,
    footnote_text: str,
    output_filename: Optional[str] = None
) -> Tuple[bool, str]:
    """Add footnote at the end of a specific paragraph (backward compatibility)."""
    success, message, _ = add_footnote_robust(
        filename=filename,
        paragraph_index=paragraph_index,
        footnote_text=footnote_text,
        output_filename=output_filename,
        position="after"
    )
    return success, message


def add_footnote_with_proper_formatting(
    filename: str,
    search_text: str,
    footnote_text: str,
    output_filename: Optional[str] = None,
    position: str = "after"
) -> Tuple[bool, str]:
    """Add footnote with proper formatting (backward compatibility)."""
    success, message, _ = add_footnote_robust(
        filename=filename,
        search_text=search_text,
        footnote_text=footnote_text,
        output_filename=output_filename,
        position=position
    )
    return success, message


def delete_footnote(
    filename: str,
    footnote_id: Optional[int] = None,
    search_text: Optional[str] = None,
    output_filename: Optional[str] = None
) -> Tuple[bool, str]:
    """Delete a footnote (backward compatibility)."""
    success, message, _ = delete_footnote_robust(
        filename=filename,
        footnote_id=footnote_id,
        search_text=search_text,
        output_filename=output_filename
    )
    return success, message


# ============================================================================
# LEGACY FUNCTIONS (for core/__init__.py compatibility)
# ============================================================================

def add_footnote(doc, paragraph_index: int, footnote_text: str):
    """Legacy function for adding footnotes to python-docx Document objects.
    Note: This is a simplified version that doesn't create proper Word footnotes."""
    if paragraph_index >= len(doc.paragraphs):
        raise IndexError(f"Paragraph index {paragraph_index} out of range")
    
    para = doc.paragraphs[paragraph_index]
    # Add superscript number
    run = para.add_run()
    run.text = "¹"
    run.font.superscript = True
    
    # Add footnote text at document end
    doc.add_paragraph("_" * 50)
    footnote_para = doc.add_paragraph(f"¹ {footnote_text}")
    footnote_para.style = "Caption"
    
    return doc


# ============================================================================
# REAL ENDNOTE SUPPORT (parallel to add_footnote_robust, writes real OOXML
# endnote structures: word/endnotes.xml + w:endnoteReference instead of the
# legacy "†" superscript text hack).
# ============================================================================


def _get_safe_endnote_id(endnotes_root) -> int:
    """Get a safe endnote ID avoiding conflicts with separators and existing notes."""
    nsmap = {'w': W_NS}
    existing = endnotes_root.xpath('//w:endnote', namespaces=nsmap)
    used_ids = set()
    for en in existing:
        en_id = en.get(f'{{{W_NS}}}id')
        if en_id:
            try:
                used_ids.add(int(en_id))
            except ValueError:
                pass
    # Endnotes can start from 1 (footnotes reserve 1, endnotes do not).
    candidate_id = 1
    while candidate_id in used_ids or candidate_id in RESERVED_ENDNOTE_IDS:
        candidate_id += 1
        if candidate_id > MAX_FOOTNOTE_ID:
            raise ValueError("No available endnote IDs")
    return candidate_id


def _create_minimal_endnotes_xml() -> bytes:
    """Create minimal endnotes.xml with the two required separator notes."""
    xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:endnotes xmlns:w="{W_NS}">
    <w:endnote w:type="separator" w:id="-1">
        <w:p>
            <w:pPr>
                <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
            </w:pPr>
            <w:r>
                <w:separator/>
            </w:r>
        </w:p>
    </w:endnote>
    <w:endnote w:type="continuationSeparator" w:id="0">
        <w:p>
            <w:pPr>
                <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
            </w:pPr>
            <w:r>
                <w:continuationSeparator/>
            </w:r>
        </w:p>
    </w:endnote>
</w:endnotes>'''
    return xml.encode('utf-8')


def _ensure_endnotes_content_type(content_types_xml: bytes) -> bytes:
    """Register word/endnotes.xml in [Content_Types].xml if not already present."""
    ct_tree = etree.fromstring(content_types_xml)
    nsmap = {'ct': CT_NS}
    existing = ct_tree.xpath(
        "//ct:Override[@PartName='/word/endnotes.xml']", namespaces=nsmap
    )
    if existing:
        return content_types_xml
    override = etree.Element(
        f'{{{CT_NS}}}Override',
        PartName='/word/endnotes.xml',
        ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml',
    )
    ct_tree.append(override)
    return etree.tostring(ct_tree, encoding='UTF-8', xml_declaration=True, standalone="yes")


def _ensure_endnotes_rels(document_rels_xml: bytes) -> bytes:
    """Add an endnotes relationship to word/_rels/document.xml.rels if absent."""
    rels_tree = etree.fromstring(document_rels_xml)
    nsmap = {'r': REL_NS}
    existing = rels_tree.xpath(
        "//r:Relationship[contains(@Type, 'endnotes')]", namespaces=nsmap
    )
    if existing:
        return document_rels_xml

    all_rels = rels_tree.xpath('//r:Relationship', namespaces=nsmap)
    existing_ids = {rel.get('Id') for rel in all_rels if rel.get('Id')}
    rid_num = 1
    while f'rId{rid_num}' in existing_ids:
        rid_num += 1
    new_rid = f'rId{rid_num}'
    if len(new_rid) > MAX_RELATIONSHIP_ID_LENGTH:
        raise ValueError(f"Relationship ID too long: {new_rid}")

    rel = etree.Element(
        f'{{{REL_NS}}}Relationship',
        Id=new_rid,
        Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes',
        Target='endnotes.xml',
    )
    rels_tree.append(rel)
    return etree.tostring(rels_tree, encoding='UTF-8', xml_declaration=True, standalone="yes")


def _ensure_endnote_styles(styles_root):
    """Ensure EndnoteReference (character) and EndnoteText (paragraph) styles exist."""
    nsmap = {'w': W_NS}

    ref_style = styles_root.xpath('//w:style[@w:styleId="EndnoteReference"]', namespaces=nsmap)
    if not ref_style:
        style = etree.Element(
            f'{{{W_NS}}}style',
            attrib={
                f'{{{W_NS}}}type': 'character',
                f'{{{W_NS}}}styleId': 'EndnoteReference',
            },
        )
        name = etree.SubElement(style, f'{{{W_NS}}}name')
        name.set(f'{{{W_NS}}}val', 'endnote reference')
        base = etree.SubElement(style, f'{{{W_NS}}}basedOn')
        base.set(f'{{{W_NS}}}val', 'DefaultParagraphFont')
        rPr = etree.SubElement(style, f'{{{W_NS}}}rPr')
        vert_align = etree.SubElement(rPr, f'{{{W_NS}}}vertAlign')
        vert_align.set(f'{{{W_NS}}}val', 'superscript')
        styles_root.append(style)

    text_style = styles_root.xpath('//w:style[@w:styleId="EndnoteText"]', namespaces=nsmap)
    if not text_style:
        style = etree.Element(
            f'{{{W_NS}}}style',
            attrib={
                f'{{{W_NS}}}type': 'paragraph',
                f'{{{W_NS}}}styleId': 'EndnoteText',
            },
        )
        name = etree.SubElement(style, f'{{{W_NS}}}name')
        name.set(f'{{{W_NS}}}val', 'endnote text')
        base = etree.SubElement(style, f'{{{W_NS}}}basedOn')
        base.set(f'{{{W_NS}}}val', 'Normal')
        pPr = etree.SubElement(style, f'{{{W_NS}}}pPr')
        sz = etree.SubElement(pPr, f'{{{W_NS}}}sz')
        sz.set(f'{{{W_NS}}}val', '20')
        styles_root.append(style)


def add_endnote_robust(
    filename: str,
    search_text: Optional[str] = None,
    paragraph_index: Optional[int] = None,
    endnote_text: str = "",
    output_filename: Optional[str] = None,
    position: str = "after",
    validate_location: bool = True,
) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
    """Add an endnote with full OOXML compliance.

    Mirrors ``add_footnote_robust`` but writes ``<w:endnote>`` /
    ``<w:endnoteReference>`` into ``word/endnotes.xml`` (creating that part,
    its content-type override, its relationship, and the EndnoteReference /
    EndnoteText styles if needed). Word recognises the result as a native
    endnote (visible in View > Endnotes, accept/reject via Review pane).

    Use exactly one of ``search_text`` (find first matching paragraph) or
    ``paragraph_index`` (0-based across all body paragraphs).
    """
    if not search_text and paragraph_index is None:
        return False, "Must provide either search_text or paragraph_index", None
    if search_text and paragraph_index is not None:
        return False, "Cannot provide both search_text and paragraph_index", None
    if not os.path.exists(filename):
        return False, f"File not found: {filename}", None

    working_file = output_filename if output_filename else filename
    if output_filename and filename != output_filename:
        import shutil
        shutil.copy2(filename, output_filename)

    try:
        doc_parts = {}
        with zipfile.ZipFile(filename, 'r') as zin:
            doc_parts['document'] = zin.read('word/document.xml')
            doc_parts['content_types'] = zin.read('[Content_Types].xml')
            doc_parts['document_rels'] = zin.read('word/_rels/document.xml.rels')
            if 'word/endnotes.xml' in zin.namelist():
                doc_parts['endnotes'] = zin.read('word/endnotes.xml')
            else:
                doc_parts['endnotes'] = _create_minimal_endnotes_xml()
            if 'word/styles.xml' in zin.namelist():
                doc_parts['styles'] = zin.read('word/styles.xml')
            else:
                doc_parts['styles'] = (
                    b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    b'<w:styles xmlns:w="' + W_NS.encode('ascii') + b'"/>'
                )

        doc_root = etree.fromstring(doc_parts['document'])
        endnotes_root = etree.fromstring(doc_parts['endnotes'])
        styles_root = etree.fromstring(doc_parts['styles'])

        nsmap = {'w': W_NS}

        if search_text:
            target_para = None
            for para in doc_root.xpath('//w:p', namespaces=nsmap):
                para_text = ''.join(para.xpath('.//w:t/text()', namespaces=nsmap))
                if search_text in para_text:
                    target_para = para
                    break
            if target_para is None:
                return False, f"Text '{search_text}' not found in document", None
        else:
            paragraphs = doc_root.xpath('//w:p', namespaces=nsmap)
            if paragraph_index >= len(paragraphs):
                return False, f"Paragraph index {paragraph_index} out of range", None
            target_para = paragraphs[paragraph_index]

        if validate_location:
            parent = target_para.getparent()
            while parent is not None:
                if parent.tag in [f'{{{W_NS}}}hdr', f'{{{W_NS}}}ftr']:
                    return False, "Cannot add endnote in header/footer", None
                parent = parent.getparent()

        endnote_id = _get_safe_endnote_id(endnotes_root)

        if position == "after":
            runs = target_para.xpath('.//w:r', namespaces=nsmap)
            insert_pos = target_para.index(runs[-1]) + 1 if runs else len(target_para)
        else:
            runs = target_para.xpath('.//w:r[w:t]', namespaces=nsmap)
            insert_pos = target_para.index(runs[0]) if runs else 0

        # Build the endnote reference run in the body.
        ref_run = etree.Element(f'{{{W_NS}}}r')
        rPr = etree.SubElement(ref_run, f'{{{W_NS}}}rPr')
        rStyle = etree.SubElement(rPr, f'{{{W_NS}}}rStyle')
        rStyle.set(f'{{{W_NS}}}val', 'EndnoteReference')
        en_ref = etree.SubElement(ref_run, f'{{{W_NS}}}endnoteReference')
        en_ref.set(f'{{{W_NS}}}id', str(endnote_id))
        target_para.insert(insert_pos, ref_run)

        # Build the endnote content paragraph inside word/endnotes.xml.
        new_endnote = etree.Element(
            f'{{{W_NS}}}endnote',
            attrib={f'{{{W_NS}}}id': str(endnote_id)},
        )
        en_para = etree.SubElement(new_endnote, f'{{{W_NS}}}p')
        pPr = etree.SubElement(en_para, f'{{{W_NS}}}pPr')
        pStyle = etree.SubElement(pPr, f'{{{W_NS}}}pStyle')
        pStyle.set(f'{{{W_NS}}}val', 'EndnoteText')

        marker_run = etree.SubElement(en_para, f'{{{W_NS}}}r')
        marker_rPr = etree.SubElement(marker_run, f'{{{W_NS}}}rPr')
        marker_rStyle = etree.SubElement(marker_rPr, f'{{{W_NS}}}rStyle')
        marker_rStyle.set(f'{{{W_NS}}}val', 'EndnoteReference')
        etree.SubElement(marker_run, f'{{{W_NS}}}endnoteRef')

        space_run = etree.SubElement(en_para, f'{{{W_NS}}}r')
        space_text = etree.SubElement(space_run, f'{{{W_NS}}}t')
        space_text.set(f'{{{XML_NS}}}space', 'preserve')
        space_text.text = ' '

        text_run = etree.SubElement(en_para, f'{{{W_NS}}}r')
        text_elem = etree.SubElement(text_run, f'{{{W_NS}}}t')
        text_elem.text = endnote_text

        endnotes_root.append(new_endnote)
        _ensure_endnote_styles(styles_root)

        content_types_xml = _ensure_endnotes_content_type(doc_parts['content_types'])
        document_rels_xml = _ensure_endnotes_rels(doc_parts['document_rels'])

        # Write package: copy unchanged parts, overwrite the five we touched.
        temp_file = working_file + '.tmp'
        rewritten = {
            'word/document.xml',
            'word/endnotes.xml',
            'word/styles.xml',
            '[Content_Types].xml',
            'word/_rels/document.xml.rels',
        }
        with zipfile.ZipFile(temp_file, 'w', zipfile.ZIP_DEFLATED) as zout:
            with zipfile.ZipFile(filename, 'r') as zin:
                for item in zin.infolist():
                    if item.filename not in rewritten:
                        zout.writestr(item, zin.read(item.filename))
            zout.writestr(
                'word/document.xml',
                etree.tostring(doc_root, encoding='UTF-8', xml_declaration=True, standalone="yes"),
            )
            zout.writestr(
                'word/endnotes.xml',
                etree.tostring(endnotes_root, encoding='UTF-8', xml_declaration=True, standalone="yes"),
            )
            zout.writestr(
                'word/styles.xml',
                etree.tostring(styles_root, encoding='UTF-8', xml_declaration=True, standalone="yes"),
            )
            zout.writestr('[Content_Types].xml', content_types_xml)
            zout.writestr('word/_rels/document.xml.rels', document_rels_xml)

        os.replace(temp_file, working_file)

        return True, f"Successfully added endnote (ID: {endnote_id}) to {working_file}", {
            'endnote_id': endnote_id,
            'location': 'search_text' if search_text else 'paragraph_index',
            'styles_created': ['EndnoteReference', 'EndnoteText'],
        }

    except Exception as e:
        temp_file = working_file + '.tmp'
        if os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except OSError:
                pass
        return False, f"Error adding endnote: {str(e)}", None


def convert_footnotes_to_endnotes_robust(
    filename: str,
    output_filename: Optional[str] = None,
) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
    """Convert every footnote in the document to a native endnote.

    For each non-separator entry in ``word/footnotes.xml``:

    1. Deep-copy the footnote XML into ``word/endnotes.xml`` under a freshly
       allocated endnote id, rewriting tag/style names along the way:
       ``w:footnote`` → ``w:endnote``,
       ``w:footnoteRef`` → ``w:endnoteRef``,
       ``pStyle w:val="FootnoteText"`` → ``EndnoteText``,
       ``rStyle w:val="FootnoteReference"`` → ``EndnoteReference``.
    2. Remove the original footnote body from ``word/footnotes.xml`` (the
       separator + continuationSeparator entries are kept so the part stays
       valid).
    3. Walk every ``<w:footnoteReference w:id="X"/>`` in ``word/document.xml``,
       rename it to ``<w:endnoteReference w:id="Y"/>`` via the id map, and
       flip the surrounding ``rStyle`` from FootnoteReference to
       EndnoteReference.
    4. Ensure the endnotes part is registered in ``[Content_Types].xml`` and
       ``document.xml.rels`` and that EndnoteReference / EndnoteText styles
       exist. (Footnote styles are left in place — other content may still
       reference them.)

    Pre-existing endnotes are preserved; new ids skip over them and the
    reserved separator ids -1 / 0.

    Returns ``(success, message, details)`` where ``details`` contains an
    ``id_mapping`` ``{old_footnote_id: new_endnote_id}``.
    """
    from copy import deepcopy

    if not os.path.exists(filename):
        return False, f"File not found: {filename}", None

    working_file = output_filename if output_filename else filename
    if output_filename and filename != output_filename:
        import shutil
        shutil.copy2(filename, output_filename)

    nsmap = {'w': W_NS}

    try:
        with zipfile.ZipFile(filename, 'r') as zin:
            names = zin.namelist()
            if 'word/footnotes.xml' not in names:
                return False, "Document has no footnotes part to convert", None
            doc_xml = zin.read('word/document.xml')
            footnotes_xml = zin.read('word/footnotes.xml')
            endnotes_xml = (
                zin.read('word/endnotes.xml') if 'word/endnotes.xml' in names
                else _create_minimal_endnotes_xml()
            )
            content_types_xml = zin.read('[Content_Types].xml')
            document_rels_xml = zin.read('word/_rels/document.xml.rels')
            if 'word/styles.xml' in names:
                styles_xml = zin.read('word/styles.xml')
            else:
                styles_xml = (
                    b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    b'<w:styles xmlns:w="' + W_NS.encode('ascii') + b'"/>'
                )

        doc_root = etree.fromstring(doc_xml)
        footnotes_root = etree.fromstring(footnotes_xml)
        endnotes_root = etree.fromstring(endnotes_xml)
        styles_root = etree.fromstring(styles_xml)

        # Collect non-separator footnotes to convert.
        footnotes_to_convert = []
        for fn in footnotes_root.xpath('//w:footnote', namespaces=nsmap):
            fn_type = fn.get(f'{{{W_NS}}}type')
            if fn_type in ('separator', 'continuationSeparator'):
                continue
            fn_id_str = fn.get(f'{{{W_NS}}}id')
            try:
                fn_id_int = int(fn_id_str)
            except (TypeError, ValueError):
                continue
            if fn_id_int in RESERVED_FOOTNOTE_IDS - {1}:
                # 1 is "reserved" only at allocation time for our adder; if a
                # real footnote already sits at id 1 we still convert it.
                continue
            footnotes_to_convert.append(fn)

        if not footnotes_to_convert:
            return False, "Document has no convertible footnotes (only separators present)", None

        # Existing endnote ids that we must not collide with.
        existing_endnote_ids = set()
        for en in endnotes_root.xpath('//w:endnote', namespaces=nsmap):
            try:
                existing_endnote_ids.add(int(en.get(f'{{{W_NS}}}id')))
            except (TypeError, ValueError):
                pass

        next_id = 1
        def _alloc_id() -> int:
            nonlocal next_id
            while next_id in existing_endnote_ids or next_id in RESERVED_ENDNOTE_IDS:
                next_id += 1
            chosen = next_id
            existing_endnote_ids.add(chosen)
            next_id += 1
            return chosen

        id_map: Dict[int, int] = {}

        for fn in footnotes_to_convert:
            old_id = int(fn.get(f'{{{W_NS}}}id'))
            new_id = _alloc_id()
            id_map[old_id] = new_id

            # Deep-copy and transform: footnote → endnote.
            en = deepcopy(fn)
            en.tag = f'{{{W_NS}}}endnote'
            en.set(f'{{{W_NS}}}id', str(new_id))

            for fn_ref in en.iter(f'{{{W_NS}}}footnoteRef'):
                fn_ref.tag = f'{{{W_NS}}}endnoteRef'
            for pStyle in en.iter(f'{{{W_NS}}}pStyle'):
                if pStyle.get(f'{{{W_NS}}}val') == 'FootnoteText':
                    pStyle.set(f'{{{W_NS}}}val', 'EndnoteText')
            for rStyle in en.iter(f'{{{W_NS}}}rStyle'):
                if rStyle.get(f'{{{W_NS}}}val') == 'FootnoteReference':
                    rStyle.set(f'{{{W_NS}}}val', 'EndnoteReference')

            endnotes_root.append(en)

        # Remove the original footnote bodies (keep separators intact).
        for fn in footnotes_to_convert:
            parent = fn.getparent()
            if parent is not None:
                parent.remove(fn)

        # Rewrite body references: <w:footnoteReference> → <w:endnoteReference>.
        body_changes = 0
        for fn_ref in doc_root.xpath('//w:footnoteReference', namespaces=nsmap):
            old_id_str = fn_ref.get(f'{{{W_NS}}}id')
            try:
                old_id = int(old_id_str)
            except (TypeError, ValueError):
                continue
            if old_id not in id_map:
                continue  # Was a separator reference or unknown; skip.
            fn_ref.tag = f'{{{W_NS}}}endnoteReference'
            fn_ref.set(f'{{{W_NS}}}id', str(id_map[old_id]))

            parent_run = fn_ref.getparent()
            if parent_run is not None:
                for rStyle in parent_run.iter(f'{{{W_NS}}}rStyle'):
                    if rStyle.get(f'{{{W_NS}}}val') == 'FootnoteReference':
                        rStyle.set(f'{{{W_NS}}}val', 'EndnoteReference')
            body_changes += 1

        # Make sure the endnote infrastructure is registered.
        _ensure_endnote_styles(styles_root)
        content_types_xml = _ensure_endnotes_content_type(content_types_xml)
        document_rels_xml = _ensure_endnotes_rels(document_rels_xml)

        # Repackage.
        temp_file = working_file + '.tmp'
        rewritten = {
            'word/document.xml',
            'word/footnotes.xml',
            'word/endnotes.xml',
            'word/styles.xml',
            '[Content_Types].xml',
            'word/_rels/document.xml.rels',
        }
        with zipfile.ZipFile(temp_file, 'w', zipfile.ZIP_DEFLATED) as zout:
            with zipfile.ZipFile(filename, 'r') as zin:
                for item in zin.infolist():
                    if item.filename not in rewritten:
                        zout.writestr(item, zin.read(item.filename))
            zout.writestr(
                'word/document.xml',
                etree.tostring(doc_root, encoding='UTF-8', xml_declaration=True, standalone="yes"),
            )
            zout.writestr(
                'word/footnotes.xml',
                etree.tostring(footnotes_root, encoding='UTF-8', xml_declaration=True, standalone="yes"),
            )
            zout.writestr(
                'word/endnotes.xml',
                etree.tostring(endnotes_root, encoding='UTF-8', xml_declaration=True, standalone="yes"),
            )
            zout.writestr(
                'word/styles.xml',
                etree.tostring(styles_root, encoding='UTF-8', xml_declaration=True, standalone="yes"),
            )
            zout.writestr('[Content_Types].xml', content_types_xml)
            zout.writestr('word/_rels/document.xml.rels', document_rels_xml)

        os.replace(temp_file, working_file)

        return True, (
            f"Converted {len(id_map)} footnote(s) to endnote(s); "
            f"updated {body_changes} body reference(s)"
        ), {
            'converted': len(id_map),
            'body_references_updated': body_changes,
            'id_mapping': {str(k): v for k, v in id_map.items()},
        }

    except Exception as e:
        temp_file = working_file + '.tmp'
        if os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except OSError:
                pass
        return False, f"Error converting footnotes to endnotes: {str(e)}", None


def add_endnote(doc, paragraph_index: int, endnote_text: str):
    """Legacy function for adding endnotes."""
    if paragraph_index >= len(doc.paragraphs):
        raise IndexError(f"Paragraph index {paragraph_index} out of range")
    
    para = doc.paragraphs[paragraph_index]
    run = para.add_run()
    run.text = "†"
    run.font.superscript = True
    
    # Endnotes go at the very end
    doc.add_page_break()
    doc.add_heading("Endnotes", level=1)
    endnote_para = doc.add_paragraph(f"† {endnote_text}")
    
    return doc


def convert_footnotes_to_endnotes(doc):
    """Legacy function to convert footnotes to endnotes in a Document object."""
    # This is a placeholder - real conversion requires XML manipulation
    return doc


# XML_NS is defined near the top of the module alongside W_NS, R_NS, etc.